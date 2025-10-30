from __future__ import annotations

import json
import os
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

from flask import Flask, abort, flash, g, jsonify, redirect, render_template, request, session, url_for
from werkzeug.security import check_password_hash, generate_password_hash

from models import db, Account, Interaction
from feedback_tech import generate_feedback
from rubric_extract import extract_rubric_from_text


ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "data"
DATA_DIR.mkdir(parents=True, exist_ok=True)
RUBRIC_PATH = DATA_DIR / "rubric_master.json"


def _load_json(path: Path, default: Any) -> Any:
    try:
        if path.exists():
            return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        pass
    return default


def _dump_json(path: Path, obj: Any) -> None:
    path.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")


DEFAULT_TECH_RUBRIC = [
    {
        "name": "Executive Summary",
        "max_points": 6,
        "scoringCriteria": [
            {"points": 6, "description": "Exemplary: Clear, concise; what, how, results, why; quantitative deltas; 1 page."},
            {"points": 4, "description": "Proficient: One quantitative element thin or minor clarity issues."},
            {"points": 2, "description": "Developing: Misses 1 required question/metric."},
            {"points": 0, "description": "Insufficient: Over length or lacks core content."}
        ],
        "notes": "Page limit: 1 page; >10% over caps at Proficient (4)."
    },
    {
        "name": "Justification (Intro & Baseline)",
        "max_points": 12,
        "scoringCriteria": [
            {"points": 12, "description": "Exemplary: All required items complete with citations/justifications; accurate PFD & complete stream table; strong problem framing."},
            {"points": 9, "description": "Proficient: Minor gaps (e.g., shallow demand forecast or limited sensitivity rationale)."},
            {"points": 6, "description": "Developing: ~2 key items missing (e.g., no thermo basis or incomplete stream table)."},
            {"points": 3, "description": "Insufficient: Major omissions (no PFD/streams or no problem statement)."}
        ],
        "gates": [
            "Missing PFD or full stream table, or missing one-to-one PFD↔simulation stream numbering → cap 7"
        ],
        "notes": "Page limit: 6 pages (stream table excluded)."
    },
    {
        "name": "Summary (Improved Process & Results)",
        "max_points": 9,
        "scoringCriteria": [
            {"points": 9, "description": "Exemplary: Complete, quantitative contrast; rationale aligns with objectives."},
            {"points": 7, "description": "Proficient: Mostly complete; one element thin (e.g., partial cost table)."},
            {"points": 5, "description": "Developing: Multiple missing pieces (e.g., no NPW/IRR or incomplete improved PFD)."},
            {"points": 2, "description": "Insufficient: Little evidence of improvement."}
        ],
        "gates": [
            "Missing improved PFD or missing NPW/IRR → cap 5",
            "Main equipment replaced without strong justification → cap 7"
        ],
        "notes": "Page limit: 6 pages."
    },
    {
        "name": "Conclusions",
        "max_points": 3,
        "scoringCriteria": [
            {"points": 3, "description": "Exemplary: Interprets results against objectives; clear go/no-go with rationale."},
            {"points": 2, "description": "Proficient: Generally sound; minor gaps."},
            {"points": 1, "description": "Developing: Re-summarizes; weak linkage to objectives."},
            {"points": 0, "description": "Insufficient."}
        ]
    },
    {
        "name": "Recommendations",
        "max_points": 3,
        "scoringCriteria": [
            {"points": 3, "description": "Exemplary: Technical/operational feasibility and economic optimum; avoids project-task items."},
            {"points": 2, "description": "Proficient: Mostly focused; minor scope creep."},
            {"points": 1, "description": "Developing: Vague or task-list oriented."},
            {"points": 0, "description": "Insufficient."}
        ]
    },
    {
        "name": "6a) Designed Equipment",
        "max_points": 18,
        "scoringCriteria": [
            {"points": 18, "description": "Exemplary: All tables/diagrams complete; numbering matches; costs/utilities/kinetics documented."},
            {"points": 15, "description": "Proficient: Minor missing attributes or thin sensitivity justification."},
            {"points": 10, "description": "Developing: Several missing tables/parameters or poor consistency."},
            {"points": 6, "description": "Insufficient: Largely incomplete."}
        ],
        "gates": [
            "No one-to-one PFD↔simulation mapping, or missing full process stream table, or missing capital-cost method → cap 12",
            "Missing Aspen backups (base & improved) → cap 12"
        ]
    },
    {
        "name": "6b) Safety, Health & Environment",
        "max_points": 9,
        "scoringCriteria": [
            {"points": 9, "description": "Exemplary: Hazards/risks/mitigations; P&ID with controls and justification; HAZOP (1 equipment); regs cited."},
            {"points": 7, "description": "Proficient: Mostly complete; minor omissions."},
            {"points": 5, "description": "Developing: Limited depth or missing elements."},
            {"points": 2, "description": "Insufficient."}
        ],
        "gates": [
            "Missing P&ID or HAZOP → cap 4"
        ]
    },
    {
        "name": "6c) Economic Analysis",
        "max_points": 9,
        "scoringCriteria": [
            {"points": 9, "description": "Exemplary: COL/CRM/CWT/CUT & COMd; cash flow and DCF; explicit assumptions; NPW & IRR; price sources."},
            {"points": 7, "description": "Proficient: Mostly complete; minor gaps."},
            {"points": 5, "description": "Developing: Missing depth or a key piece."},
            {"points": 2, "description": "Insufficient."}
        ],
        "gates": [
            "Missing DCF or NPW/IRR → cap 4"
        ]
    },
    {
        "name": "Quality of the Written Work",
        "max_points": 6,
        "scoringCriteria": [
            {"points": 6, "description": "Exemplary: Fully compliant; professional polish; citation and numbering rules followed."},
            {"points": 4, "description": "Proficient: Minor inconsistencies."},
            {"points": 2, "description": "Developing: Multiple format or clarity issues."},
            {"points": 1, "description": "Insufficient: Non-compliant or hard to read."}
        ],
        "notes": "Formatting rules: Word; two-sided justification; 8.5×11; body 1.15; Arial 12 (titles 10); each major section new page; prelim pages Roman numerals; chapter-scoped numbering; in-text references ordered with matching bibliography."
    }
]

if not RUBRIC_PATH.exists():
    _dump_json(RUBRIC_PATH, DEFAULT_TECH_RUBRIC)


def create_app() -> Flask:
    app = Flask(__name__)
    app.config.update(
        SECRET_KEY=os.getenv("SECRET_KEY", "dev-secret"),
        JSON_AS_ASCII=False,
        PERMANENT_SESSION_LIFETIME=timedelta(days=14),
    )

    db_url = os.getenv("DATABASE_URL")
    if db_url and db_url.startswith("postgres://"):
        db_url = db_url.replace("postgres://", "postgresql+psycopg2://", 1)

    # default to SQLite for local dev if DATABASE_URL is not provided
    instance_dir = ROOT / "instance"
    instance_dir.mkdir(exist_ok=True)
    sqlite_path = instance_dir / "app.db"

    app.config.update(
        SQLALCHEMY_DATABASE_URI=db_url or f"sqlite:///{sqlite_path}",
        SQLALCHEMY_TRACK_MODIFICATIONS=False,
        SQLALCHEMY_ENGINE_OPTIONS={"pool_pre_ping": True, "pool_recycle": 300},
        SESSION_COOKIE_HTTPONLY=True,
        SESSION_COOKIE_SAMESITE="Lax",
        SESSION_COOKIE_SECURE=True,
    )

    db.init_app(app)

    with app.app_context():
        db.create_all()
        _maybe_seed_admin()

    # ----- user/session helpers -----
    @app.before_request
    def _load_current_user():
        g.current_user = None
        uid = session.get("user_id")
        if uid:
            g.current_user = db.session.get(Account, uid)

    @app.context_processor
    def _inject_tpl_vars():
        cu = getattr(g, "current_user", None)
        return {
            "current_user": cu,
            "is_admin": bool(cu and getattr(cu, "role", "") == "admin"),
        }

    # ----- auth pages -----
    @app.route("/login", methods=["GET", "POST"])
    def login():
        if "user_id" in session and db.session.get(Account, session["user_id"]):
            return redirect(url_for("index"))
        if request.method == "POST":
            username = request.form.get("username", "").strip()
            password = request.form.get("password", "")
            acc = Account.query.filter_by(username=username).first()
            if acc is None or not check_password_hash(acc.password_hash, password):
                flash("Invalid username or password.", "danger")
                return redirect(url_for("login"))
            session["user_id"] = acc.id
            session.permanent = True
            return redirect(url_for("index"))
        return render_template("login.html")

    @app.route("/logout")
    def logout():
        session.pop("user_id", None)
        flash("You have been logged out.", "success")
        return redirect(url_for("login"))

    @app.route("/register", methods=["GET", "POST"])
    def register():
        if request.method == "POST":
            username = request.form.get("username", "").strip()
            password = request.form.get("password", "")
            confirm = request.form.get("confirm_password", "")
            if not username or not password:
                flash("Username and password are required.", "danger")
                return redirect(url_for("register"))
            if password != confirm:
                flash("Passwords do not match.", "danger")
                return redirect(url_for("register"))
            if Account.query.filter_by(username=username).first():
                flash("Username already exists.", "danger")
                return redirect(url_for("register"))
            acc = Account(username=username, password_hash=generate_password_hash(password))
            db.session.add(acc)
            db.session.commit()
            flash("Registration successful. Please log in.", "success")
            return redirect(url_for("login"))
        return render_template("register.html")

    # ----- pages -----
    @app.route("/")
    def index():
        if not getattr(g, "current_user", None):
            return redirect(url_for("login"))
        rubric = _load_json(RUBRIC_PATH, DEFAULT_TECH_RUBRIC)
        return render_template("index.html", rubric=rubric)

    # ----- rubric endpoints (compat with prior app) -----
    @app.get("/get_WRITING_RUBRICs")
    def get_rubrics():
        rubric = _load_json(RUBRIC_PATH, DEFAULT_TECH_RUBRIC)
        return jsonify(rubric)

    def _admin_required() -> None:
        cu = getattr(g, "current_user", None)
        if not cu or getattr(cu, "role", "") != "admin":
            abort(403)

    @app.post("/save_WRITING_RUBRICs")
    def save_rubrics():
        _admin_required()
        try:
            data = request.get_json(force=True)
            if not isinstance(data, list):
                return jsonify({"error": "Body must be a JSON array of rubrics"}), 400
            # Save to file
            _dump_json(RUBRIC_PATH, data)
            # Versioning record
            try:
                from models import RubricVersion
                rv = RubricVersion(created_by=getattr(g.current_user, 'id', None), rubric_json=json.dumps(data, ensure_ascii=False))
                db.session.add(rv)
                db.session.commit()
            except Exception:
                db.session.rollback()
            return jsonify({"message": "Rubrics saved"})
        except Exception as e:
            return jsonify({"error": str(e)}), 400

    @app.post("/rubric/extract")
    def rubric_extract():
        """Admin-only: extract a rubric from an uploaded syllabus (PDF/DOCX/TXT)."""
        _admin_required()
        f = request.files.get("file")
        if not f or not f.filename:
            return jsonify({"success": False, "error": "Please upload a syllabus file (PDF/DOCX/TXT)."}), 400
        fname = f.filename.lower()
        text = ""
        try:
            if fname.endswith(".pdf"):
                from pdfminer.high_level import extract_text as _pdf_extract
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                    f.save(tmp.name)
                    text = _pdf_extract(tmp.name) or ""
            elif fname.endswith(".docx"):
                import docx as _docx
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    f.save(tmp.name)
                    doc = _docx.Document(tmp.name)
                    text = "\n".join(p.text for p in doc.paragraphs)
            else:
                text = f.read().decode("utf-8", "ignore")
        except Exception:
            text = text or ""

        text = (text or "").strip()
        if not text:
            return jsonify({"success": False, "error": "Could not read text from the uploaded file."}), 400

        rubric = extract_rubric_from_text(text)
        if not isinstance(rubric, list) or not rubric:
            return jsonify({"success": False, "error": "No rubric could be extracted."}), 200
        return jsonify({"success": True, "rubric": rubric})

    @app.get("/rubric/versions")
    def rubric_versions():
        _admin_required()
        from models import RubricVersion
        rows = (RubricVersion.query.order_by(RubricVersion.created_at.desc()).limit(20).all())
        return jsonify({
            "success": True,
            "versions": [
                {
                    "id": r.id,
                    "created_at": r.created_at.isoformat(),
                    "created_by": r.created_by,
                } for r in rows
            ]
        })

    @app.post("/rubric/rollback")
    def rubric_rollback():
        _admin_required()
        try:
            body = request.get_json(force=True) or {}
            vid = int(body.get("version_id"))
        except Exception:
            return jsonify({"success": False, "error": "version_id is required"}), 400
        from models import RubricVersion
        r = db.session.get(RubricVersion, vid)
        if not r:
            return jsonify({"success": False, "error": "Version not found"}), 404
        try:
            obj = json.loads(r.rubric_json)
        except Exception:
            return jsonify({"success": False, "error": "Stored version JSON invalid"}), 400
        _dump_json(RUBRIC_PATH, obj)
        # Also append a new version entry noting rollback
        try:
            rv = RubricVersion(created_by=getattr(g.current_user, 'id', None), rubric_json=r.rubric_json)
            db.session.add(rv)
            db.session.commit()
        except Exception:
            db.session.rollback()
        return jsonify({"success": True, "message": "Rolled back to selected version"})

    @app.get("/history")
    def history_list():
        if not getattr(g, "current_user", None):
            abort(401)
        rows = (Interaction.query
                .filter(Interaction.user_id == g.current_user.id, Interaction.feedback_text.isnot(None))
                .order_by(Interaction.feedback_time.desc())
                .limit(30)
                .all())
        def excerpt(s: str, n: int = 160) -> str:
            s = (s or "")
            return (s[:n] + ("..." if len(s) > n else "")) if s else ""
        return jsonify({
            "success": True,
            "items": [
                {
                    "id": r.id,
                    "feedback_time": (r.feedback_time.isoformat() if r.feedback_time else None),
                    "prompt_excerpt": excerpt(r.prompt_text or "", 160),
                    "feedback_excerpt": excerpt(r.feedback_text or "", 160),
                } for r in rows
            ]
        })

    @app.get("/history/<int:rid>")
    def history_detail(rid: int):
        if not getattr(g, "current_user", None):
            abort(401)
        r = db.session.get(Interaction, rid)
        if (not r) or (r.user_id != g.current_user.id):
            abort(404)
        return jsonify({
            "success": True,
            "id": r.id,
            "feedback_time": (r.feedback_time.isoformat() if r.feedback_time else None),
            "prompt_text": r.prompt_text or "",
            "feedback_text": r.feedback_text or "",
            "scores": json.loads(r.scores_json or "{}") if r.scores_json else {},
        })

    @app.get("/export_pdf")
    def export_pdf():
        if not getattr(g, "current_user", None):
            abort(401)
        try:
            rid = int(request.args.get("interaction_id", "0"))
        except Exception:
            return jsonify({"success": False, "error": "interaction_id is required"}), 400
        r = db.session.get(Interaction, rid)
        if (not r) or (r.user_id != g.current_user.id):
            abort(404)
        # Build a simple PDF (wrap long tokens to avoid width errors)
        from fpdf import FPDF
        pdf = FPDF()
        pdf.set_margins(12, 12, 12)
        pdf.set_auto_page_break(auto=True, margin=12)
        pdf.add_page()
        pdf.set_font('Helvetica', size=12)
        # Helper to wrap very long tokens with soft breaks
        def _wrap_for_pdf(s: str, max_token: int = 50) -> str:
            if not s:
                return ""
            out_lines = []
            for line in s.split('\n'):
                parts = line.split(' ')
                fixed_parts = []
                for p in parts:
                    if len(p) > max_token:
                        # insert spaces every max_token chars to allow wrapping
                        chunks = [p[i:i+max_token] for i in range(0, len(p), max_token)]
                        fixed_parts.append(' '.join(chunks))
                    else:
                        fixed_parts.append(p)
                out_lines.append(' '.join(fixed_parts))
            return '\n'.join(out_lines)

        # Replace non-Latin-1 punctuation and drop unsupported chars for core fonts
        def _latin1_safe(s: str) -> str:
            if not s:
                return ""
            repl = {
                '\u2013': '-', '\u2014': '-',  # en/em dash
                '\u2018': "'", '\u2019': "'",  # single quotes
                '\u201C': '"', '\u201D': '"',  # double quotes
                '\u2026': '...',               # ellipsis
                '\u00A0': ' ',                 # nbsp
            }
            for k, v in repl.items():
                s = s.replace(k, v)
            # keep only latin-1 range
            return s.encode('latin-1', 'ignore').decode('latin-1')
        title = "Technical Report Feedback"
        pdf.set_text_color(20, 20, 20)
        epw = pdf.w - pdf.l_margin - pdf.r_margin
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 8, _latin1_safe(title), align='L')
        pdf.set_text_color(80, 80, 80)
        ts = r.feedback_time.isoformat() if r.feedback_time else ""
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, _latin1_safe(f"Time: {ts}"), align='L')
        pdf.ln(2)
        pdf.set_text_color(20, 20, 20)
        txt = r.feedback_text or ""
        if not (txt.strip()):
            hint = "No feedback available yet."
            if not os.getenv("OPENAI_API_KEY", "").strip():
                hint += " LLM offline (no OPENAI_API_KEY)."
            txt = hint
        from fpdf.errors import FPDFException
        for para in (txt.split("\n") if txt else []):
            ptext = _latin1_safe(_wrap_for_pdf(para))
            try:
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 6, ptext, align='L')
            except FPDFException:
                # Fallback: switch to Courier and try again with smaller size
                pdf.set_font('Courier', size=11)
                try:
                    pdf.set_x(pdf.l_margin)
                    pdf.multi_cell(0, 6, ptext, align='L')
                except FPDFException:
                    # Final fallback: insert soft spaces between characters
                    tight = ' '.join(list(ptext))
                    pdf.set_x(pdf.l_margin)
                    pdf.multi_cell(0, 6, tight, align='L')
                finally:
                    pdf.set_font('Helvetica', size=12)
        out = pdf.output(dest='S')
        if isinstance(out, str):
            out = out.encode('latin1', 'ignore')
        import io
        from flask import send_file
        return send_file(
            io.BytesIO(out),
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'feedback_{r.id}.pdf'
        )

    # ----- feedback endpoints -----
    @app.get("/get_last_feedback")
    def get_last_feedback():
        if not getattr(g, "current_user", None):
            abort(401)
        rec = (
            Interaction.query
            .filter(Interaction.user_id == g.current_user.id, Interaction.feedback_text.isnot(None))
            .order_by(Interaction.feedback_time.desc())
            .first()
        )
        return jsonify({"success": True, "feedback": rec.feedback_text if rec else ""})

    @app.post("/get_feedback")
    def post_get_feedback():
        if not getattr(g, "current_user", None):
            abort(401)

        msg = request.form.get("message", "").strip()
        f = request.files.get("file")

        uploaded_text = ""
        if f and f.filename:
            fname = f.filename.lower()
            try:
                if fname.endswith(".pdf"):
                    from pdfminer.high_level import extract_text as _pdf_extract
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                        f.save(tmp.name)
                        uploaded_text = _pdf_extract(tmp.name) or ""
                elif fname.endswith(".docx"):
                    import docx as _docx
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                        f.save(tmp.name)
                        doc = _docx.Document(tmp.name)
                        uploaded_text = "\n".join(p.text for p in doc.paragraphs)
                else:
                    uploaded_text = f.read().decode("utf-8", "ignore")
            except Exception:
                uploaded_text = uploaded_text or ""

        narrative_text = msg or uploaded_text

        rubric = _load_json(RUBRIC_PATH, DEFAULT_TECH_RUBRIC)
        try:
            feedback_text, scores, feedback_summary, evidence_quotes = generate_feedback(
                message=narrative_text,
                rubric=rubric,
            )
        except Exception:
            # Backward-compat if function returns only 3 values
            ft, sc, sm = generate_feedback(message=narrative_text, rubric=rubric)  # type: ignore
            feedback_text, scores, feedback_summary = ft, sc, sm
            evidence_quotes = []

        rec = Interaction(
            user_id=g.current_user.id,
            prompt_text=narrative_text or None,
            prompt_time=datetime.now(timezone.utc),
            feedback_text=feedback_text,
            feedback_summary=feedback_summary,
            feedback_time=datetime.now(timezone.utc),
            scores_json=json.dumps(scores, ensure_ascii=False),
            status="final",
        )
        db.session.add(rec)
        db.session.commit()

        return jsonify({
            "success": True,
            "feedback": feedback_text,
            "feedback_summary": feedback_summary,
            "scores": scores,
            "evidence_quotes": evidence_quotes,
            "interaction_id": rec.id,
            "prompt_excerpt": (narrative_text or "")[:4000],
        })

    @app.post("/submit_feedback")
    def submit_feedback():
        if not getattr(g, "current_user", None):
            abort(401)
        try:
            data = request.get_json(force=True) or {}
            interaction_id = data.get("interaction_id")
            rating = data.get("rating")
            feedback = data.get("feedback")

            rec = None
            if interaction_id:
                rec = db.session.get(Interaction, interaction_id)
            if not rec:
                # fallback to latest
                rec = (
                    Interaction.query
                    .filter(Interaction.user_id == g.current_user.id)
                    .order_by(Interaction.feedback_time.desc())
                    .first()
                )
            if not rec:
                return jsonify({"success": False, "error": "record not found"}), 404

            if rating is not None:
                try:
                    rec.rating = int(rating)
                except Exception:
                    pass
            if feedback is not None:
                rec.student_feedback_text = str(feedback)

            db.session.commit()
            return jsonify({"success": True})
        except Exception as e:
            db.session.rollback()
            return jsonify({"success": False, "error": str(e)}), 400

    return app


def _maybe_seed_admin() -> None:
    username = os.getenv("ADMIN_USERNAME", "").strip()
    password = os.getenv("ADMIN_PASSWORD", "").strip()
    if not username or not password:
        return
    exists = Account.query.filter_by(username=username).first()
    if exists:
        return
    acc = Account(username=username, role="admin", password_hash=generate_password_hash(password))
    db.session.add(acc)
    db.session.commit()


app = create_app()

if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5000, debug=True)
